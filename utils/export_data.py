from docx import Document
from io import BytesIO
import matplotlib.pyplot as plt

def export_to_work(results):
    doc= Document()
    doc.add_heading("Data Analysis Report",0)

    for fig, text in results:
        try:
            #save figures to aa bytes buffer
            img_buffer = BytesIO()
            fig.savefig(img_buffer,format='PNG',dpi=300, bbox_inches='tight')
            img_buffer.seek(0)

            #add to doc
            doc.add_picture(img_buffer)
            doc.add_paragraph(text)
            doc.add_paragraph("") # la nerja3 3al sater

            #close the figure
            plt.close(fig)
        
        except Exception as e:
            print(f"Error adding figure to document: {e}")
            continue
    
    #save doc
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output

    